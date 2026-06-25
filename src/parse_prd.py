from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document

from utils import ensure_dir, write_json


KEYWORDS = {
    "background": ["背景", "目标", "范围"],
    "roles": ["角色", "权限", "用户"],
    "process": ["流程", "节点", "状态", "审批"],
    "rules": ["规则", "触发", "限制", "校验"],
    "pages": ["页面", "PC", "移动", "小程序", "菜单"],
    "faq": ["问题", "待确认", "FAQ"],
}


def read_docx(path: Path) -> list[str]:
    doc = Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return paragraphs


def read_markdown(path: Path) -> list[str]:
    return [line.strip() for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


def bucket_lines(lines: list[str]) -> dict[str, list[str]]:
    buckets = {key: [] for key in KEYWORDS}
    buckets["other"] = []
    for line in lines:
        matched = False
        for key, words in KEYWORDS.items():
            if any(word.lower() in line.lower() for word in words):
                buckets[key].append(line)
                matched = True
        if not matched:
            buckets["other"].append(line)
    return buckets


def parse_prd(prd_path: Path) -> dict:
    if prd_path.suffix.lower() == ".docx":
        lines = read_docx(prd_path)
    else:
        lines = read_markdown(prd_path)
    title = next((line.lstrip("# ").strip() for line in lines if line.strip("# ").strip()), "示例产品操作手册")
    return {
        "source": prd_path.name,
        "title": title,
        "paragraph_count": len(lines),
        "lines": lines,
        "sections": bucket_lines(lines),
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Parse PRD docx/md into structured JSON.")
    parser.add_argument("prd", type=Path)
    parser.add_argument("--out", type=Path, default=Path("plans/parsed_prd.json"))
    args = parser.parse_args()

    data = parse_prd(args.prd)
    ensure_dir(args.out.parent)
    write_json(args.out, data)
    print(f"parsed: {args.prd} -> {args.out}")


if __name__ == "__main__":
    main()
